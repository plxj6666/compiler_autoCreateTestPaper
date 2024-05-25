import tkinter as tk
from tkinter import ttk, messagebox
import random
import ast
from fpdf import FPDF
from docx import Document, shared
import os

# 读取题库数据
def read_questions(file_path):
    questions = []
    with open(file_path, 'r', encoding='utf-8') as file:
        for line in file:
            if line.strip():  # 跳过空行
                try:
                    question = ast.literal_eval(line.strip())
                    questions.append(question)
                except (SyntaxError, ValueError) as e:
                    print(f"Error parsing line: {line}. Error: {e}")
    return questions

# 统计每种题型的题量
def count_questions(questions):
    counts = {'单选题': 0, '多选题': 0, '判断题': 0, '简答题': 0}
    for question in questions:
        qtype = question[0].strip("（").strip()  # 移除可能存在的括号
        if qtype in counts:
            counts[qtype] += 1
        else:
            print(f"Unrecognized question type: {qtype}")
    return counts

# 创建PDF文档
def create_pdf(exam_name, content):
    pdf = FPDF()
    pdf.add_page()
    font_path = os.path.join(os.path.dirname(__file__), './src/SimSun.ttf')
    pdf.add_font('SimSun', '', font_path, uni=True)
    pdf.set_font('SimSun', '', 16)
    pdf.cell(0, 10, exam_name, 0, 1, 'C')
    pdf.set_font('SimSun', '', 12)
    pdf.multi_cell(0, 10, content)
    if not os.path.exists("./output"):
        os.makedirs("./output")
    pdf.output(f"./output/{exam_name}.pdf")

# 创建Word文档
def create_word(exam_name, content):
    doc = Document()
    doc.add_heading(exam_name, level=1)
    for line in content.split('\n'):
        if line.strip():
            p = doc.add_paragraph(line)
            if line.startswith("第") and "部分" in line:
                p.bold = True
            elif line.startswith("   "):
                p.left_indent = shared.Inches(0.5)
    if not os.path.exists("./output"):
        os.makedirs("./output")
    doc.save(f"./output/{exam_name}.docx")

# 创建GUI
class ExamGenerator(tk.Tk):
    def __init__(self, questions):
        super().__init__()
        self.title("试卷生成器")
        self.questions = questions
        self.counts = count_questions(questions)

        # 试卷名称输入
        self.exam_name_label = tk.Label(self, text="试卷名称")
        self.exam_name_label.grid(row=0, column=0, padx=10, pady=10)
        self.exam_name_entry = tk.Entry(self)
        self.exam_name_entry.grid(row=0, column=1, padx=10, pady=10)

        # 每种题型的选择数量
        self.selections = {}
        row = 1
        for qtype in self.counts.keys():
            label = tk.Label(self, text=f"{qtype}题数量（最多{self.counts[qtype]}题）")
            label.grid(row=row, column=0, padx=10, pady=10)
            combobox = ttk.Combobox(self, values=list(range(self.counts[qtype] + 1)))
            combobox.current(0)
            combobox.grid(row=row, column=1, padx=10, pady=10)
            self.selections[qtype] = combobox
            row += 1

        # 生成试卷按钮
        self.generate_button = tk.Button(self, text="生成试卷", command=self.generate_exam)
        self.generate_button.grid(row=row, column=0, columnspan=2, padx=10, pady=10)

        # 试卷显示区域
        self.exam_text = tk.Text(self, width=80, height=20)
        self.exam_text.grid(row=row + 1, column=0, columnspan=2, padx=10, pady=10)

    def generate_exam(self):
        exam_name = self.exam_name_entry.get()
        if not exam_name:
            messagebox.showerror("错误", "请填写试卷名称")
            return

        selected_counts = {qtype: int(combobox.get()) for qtype, combobox in self.selections.items()}
        exam_content = self.create_exam_content(exam_name, selected_counts)
        self.exam_text.delete(1.0, tk.END)
        self.exam_text.insert(tk.END, exam_content)

        # 创建PDF和Word文件
        create_pdf(exam_name, exam_content)
        create_word(exam_name, exam_content)
        messagebox.showinfo("成功", f"试卷已生成为PDF和Word文件：./output/{exam_name}.pdf, ./output/{exam_name}.docx")

    def create_exam_content(self, exam_name, selected_counts):
        exam_content = f"{exam_name}\n"
        exam_content += "=============================================================\n"

        question_pool = {qtype: [q for q in self.questions if q[0].strip("（").strip() == qtype] for qtype in selected_counts.keys()}
        section_number = 1

        for qtype, count in selected_counts.items():
            if count == 0:
                continue
            total_score = sum(int(q[2].replace('分', '')) for q in question_pool[qtype][:count])
            exam_content += f"第{section_number}部分 {qtype}（共{total_score}分，{count}题）\n"
            section_number += 1
            selected_questions = random.sample(question_pool[qtype], count)
            for idx, q in enumerate(selected_questions, start=1):
                exam_content += f"{idx}、（{q[1]}）{q[3]}（{q[2]}）\n"
                options = q[4:]
                for option in options:
                    exam_content += f"   {option}\n"
            exam_content += "\n"
        return exam_content

if __name__ == "__main__":
    file_path = "./output/QuestionBank.txt"  # 请确保文件路径正确
    questions = read_questions(file_path)
    app = ExamGenerator(questions)
    app.mainloop()
